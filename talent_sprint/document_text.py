"""
Shared document text extraction.

Both the resume parser and the company document parser route every file through
extract_text(), which dispatches on the file extension and returns clean text.
Supported formats: pdf, docx, doc, rtf, pages, txt. (No OCR: a scanned, image
only PDF yields empty text.)

Design notes per format:
  - pdf   : pdfplumber text layer
  - docx  : python-docx (paragraphs and table cells)
  - rtf   : striprtf (pure python)
  - txt   : plain read, tolerant decoding
  - doc   : legacy binary Word has no clean pure python reader, so we shell out
            to whichever converter is on the system (macOS textutil, then
            LibreOffice, antiword, or catdoc). If none is present we warn and
            return empty text rather than crash.
  - pages : an Apple Pages file is a zip bundle; modern ones embed a QuickLook
            PDF preview, which we extract and read as a pdf. If there is no
            preview we warn and return empty.

Nothing here raises on a bad file; failures are logged and become empty text,
matching the rest of the pipeline's defensive style.
"""

import io
import os
import re
import shutil
import subprocess
import zipfile

import pdfplumber

# Every extension the pipeline will attempt to read.
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".rtf", ".pages", ".txt"}


def _clean_text(raw):
    """Light cleanup only: collapse whitespace and blank lines, then strip."""
    if not raw:
        return ""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)   # collapse spaces and tabs
    text = re.sub(r" *\n", "\n", text)    # trim trailing spaces per line
    text = re.sub(r"\n{3,}", "\n\n", text)  # collapse 3+ blank lines
    return text.strip()


def _from_pdf_stream(stream):
    """Extract text from a pdf given a path or a file-like object."""
    pages = []
    with pdfplumber.open(stream) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _from_pdf(path):
    return _from_pdf_stream(path)


def _from_docx(path):
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # Include table cell text, which resumes often use for layout.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _from_rtf(path):
    from striprtf.striprtf import rtf_to_text

    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return rtf_to_text(f.read())


def _from_txt(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _from_doc(path):
    """Legacy .doc via the first available system converter."""
    # macOS textutil: write plain text to stdout.
    if shutil.which("textutil"):
        out = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", path],
            capture_output=True, text=True, check=False,
        )
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout
    # antiword / catdoc: print text to stdout.
    for tool in ("antiword", "catdoc"):
        if shutil.which(tool):
            out = subprocess.run([tool, path], capture_output=True, text=True, check=False)
            if out.returncode == 0 and out.stdout.strip():
                return out.stdout
    # LibreOffice: convert to a temp .txt then read it.
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        import tempfile
        with tempfile.TemporaryDirectory() as tmp:
            out = subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmp, path],
                capture_output=True, text=True, check=False,
            )
            txt_path = os.path.join(tmp, os.path.splitext(os.path.basename(path))[0] + ".txt")
            if os.path.exists(txt_path):
                with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()
    print(
        f"NOTE: cannot read legacy .doc (no textutil/antiword/catdoc/libreoffice): {path}"
    )
    return ""


def _from_pages(path):
    """Apple Pages bundle: read the embedded QuickLook PDF preview if present."""
    try:
        with zipfile.ZipFile(path) as z:
            names = z.namelist()
            previews = [n for n in names if n.lower().endswith("preview.pdf")]
            if not previews:
                previews = [n for n in names if n.lower().endswith(".pdf")]
            if previews:
                return _from_pdf_stream(io.BytesIO(z.read(previews[0])))
    except Exception as exc:  # noqa: BLE001 - bad bundle must not crash the run
        print(f"WARNING: could not read .pages bundle {path}: {exc}")
        return ""
    print(f"NOTE: .pages file has no embedded PDF preview to read: {path}")
    return ""


_DISPATCH = {
    ".pdf": _from_pdf,
    ".docx": _from_docx,
    ".doc": _from_doc,
    ".rtf": _from_rtf,
    ".pages": _from_pages,
    ".txt": _from_txt,
}


def extract_text(path):
    """
    Extract clean text from a document of any supported format.

    Returns cleaned text, or an empty string for an unsupported or unreadable
    file (a warning is printed). Never raises.
    """
    ext = os.path.splitext(path)[1].lower()
    reader = _DISPATCH.get(ext)
    if reader is None:
        print(f"NOTE: unsupported file format, skipping: {path}")
        return ""
    try:
        return _clean_text(reader(path))
    except Exception as exc:  # noqa: BLE001 - one bad file must not stop the run
        print(f"WARNING: failed to extract text from {path}: {exc}")
        return ""
